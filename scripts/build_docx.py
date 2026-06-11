"""Convert the manuscript / cover letter Markdown to submission-ready .docx.

Tailored to our files' Markdown subset: H1-H3, paragraphs, **bold**, *italic*,
`code`, > blockquotes, - and 1. lists, pipe tables, --- rules, and figure
embedding (an image is inserted above any caption line that references a
figures/*.png path). Not a general Markdown engine -- just enough for these docs.
"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

_INLINE = re.compile(r"(\*\*.+?\*\*|`.+?`|\*.+?\*)")


def add_runs(paragraph, text):
    """Add inline-formatted runs (**bold**, *italic*, `code`) to a paragraph."""
    # strip markdown links -> keep text + (url)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = paragraph.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = paragraph.add_run(tok[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = paragraph.add_run(tok[1:-1]); r.italic = True
        else:
            paragraph.add_run(tok)


def emit_table(doc, rows):
    """rows: list of lists of cell strings (header first, separator already removed)."""
    if not rows:
        return
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(ncol):
            val = row[j] if j < len(row) else ""
            cells[j].paragraphs[0].text = ""
            add_runs(cells[j].paragraphs[0], val)
            if i == 0:
                for run in cells[j].paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()


def convert(md_path, out_path, embed_figures=True):
    with open(md_path, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    base = doc.styles["Normal"]
    base.font.name = "Calibri"; base.font.size = Pt(11)

    table_buf = []
    i = 0

    def flush_table():
        nonlocal table_buf
        if table_buf:
            # drop the |---| separator row
            rows = [r for r in table_buf if not re.match(r"^\s*\|?[\s:|-]+\|?\s*$", r)]
            parsed = [[c.strip() for c in re.split(r"(?<!\\)\|", r.strip().strip("|"))] for r in rows]
            emit_table(doc, parsed)
            table_buf = []

    for raw in lines:
        line = raw.rstrip("\n")
        s = line.strip()

        if s.startswith("|") and s.endswith("|"):
            table_buf.append(s); continue
        else:
            flush_table()

        if not s:
            continue
        if s == "---":
            continue

        # figure embedding: caption line referencing a figures/*.png
        m = re.search(r"figures/([A-Za-z0-9_]+\.png)", s)
        if embed_figures and m:
            img = os.path.join(ROOT, "figures", m.group(1))
            if os.path.exists(img):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    p.add_run().add_picture(img, width=Inches(6.0))
                except Exception:
                    pass
            cap = doc.add_paragraph()
            cap.paragraph_format.space_after = Pt(10)
            add_runs(cap, re.sub(r"^[-*]\s+", "", s))
            for r in cap.runs:
                r.font.size = Pt(9.5)
            continue

        if s.startswith("# "):
            h = doc.add_heading(level=0); add_runs(h, s[2:])
        elif s.startswith("## "):
            h = doc.add_heading(level=1); add_runs(h, s[3:])
        elif s.startswith("### "):
            h = doc.add_heading(level=2); add_runs(h, s[4:])
        elif s.startswith("> "):
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.4)
            add_runs(p, s[2:]);
            for r in p.runs: r.italic = True
        elif re.match(r"^[-*]\s+", s):
            p = doc.add_paragraph(style="List Bullet"); add_runs(p, re.sub(r"^[-*]\s+", "", s))
        elif re.match(r"^\d+\.\s+", s):
            p = doc.add_paragraph(style="List Number"); add_runs(p, re.sub(r"^\d+\.\s+", "", s))
        else:
            p = doc.add_paragraph(); add_runs(p, s)

    flush_table()
    doc.save(out_path)
    print("wrote", out_path)


if __name__ == "__main__":
    outdir = os.path.join(ROOT, "docs", "manuscript", "submission")
    os.makedirs(outdir, exist_ok=True)
    convert(os.path.join(ROOT, "docs", "manuscript", "manuscript.md"),
            os.path.join(outdir, "manuscript.docx"), embed_figures=True)
    convert(os.path.join(ROOT, "docs", "manuscript", "cover_letter.md"),
            os.path.join(outdir, "cover_letter.docx"), embed_figures=False)
